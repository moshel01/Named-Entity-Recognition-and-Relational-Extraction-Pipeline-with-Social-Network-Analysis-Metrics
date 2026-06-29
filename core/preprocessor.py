# Ingest files/URLs/text -> normalized plaintext. Fail-soft per source.

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Iterator, Optional

from .schema import Document, stable_id

logger = logging.getLogger(__name__)

# Extensions handled by dedicated extractors; everything else -> plaintext read.
_BINARY_LIKE = {".pdf", ".docx", ".rtf", ".html", ".htm"}
_SKIP_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".zip", ".gz", ".exe",
                  ".bin", ".mp3", ".mp4", ".mov", ".xlsx", ".pptx"}


def _read_bytes(path: Path) -> bytes:
    return path.read_bytes()


def _decode(raw: bytes, encoding: str) -> str:
    if encoding and encoding != "auto":
        return raw.decode(encoding, errors="replace")
    try:
        import chardet
        guess = chardet.detect(raw[:200_000]) or {}
        enc = guess.get("encoding") or "utf-8"
    except Exception:
        enc = "utf-8"
    return raw.decode(enc, errors="replace")


def _extract_pdf(path: Path) -> str:
    import fitz  # PyMuPDF
    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


def _extract_rtf(path: Path, encoding: str) -> str:
    from striprtf.striprtf import rtf_to_text
    raw = _decode(_read_bytes(path), encoding)
    return rtf_to_text(raw)


# Section headings that begin reference/citation/navigation tails. A bibliography
# is full of publishers and cited-author names - real entities, but not actors in
# the page's social network. Cutting the tail at the source beats tagging it after
# NER has already lifted "Oxford University Press" into a top-mention node. German
# included for de.wikipedia. The whole-word/line match + back-half + half-length
# guards keep a legitimate mid-document "Notes" from truncating the body.
_REF_HEADINGS = {
    "references", "reference", "bibliography", "notes", "footnotes", "sources",
    "works cited", "further reading", "external links", "citations", "see also",
    "literature", "literatur", "weblinks", "einzelnachweise", "quellen",
}


def _strip_trailing_sections(text: str) -> str:
    lines = text.split("\n")
    n = len(lines)
    if n < 10:
        return text
    cut = None
    for i in range(n // 2, n):       # only the back half can be a reference tail
        ln = lines[i].strip().lstrip("#").strip().rstrip(":").strip().lower()
        if ln in _REF_HEADINGS:
            cut = i
            break
    if cut is None:
        return text
    kept = "\n".join(lines[:cut]).strip()
    if len(kept) < 0.5 * len(text):  # never nuke more than half the document
        return text
    return kept


def _clean_html(raw: str) -> str:
    """HTML -> readable main text.

    Prefer trafilatura (main-content extraction: drops nav/ads/sidebars/related
    links/comments, keeps the article body + data tables) - a raw BeautifulSoup
    get_text dumps all the boilerplate too, which feeds noise into NER/RE on
    scraped pages. trafilatura is optional (`pip install trafilatura`); a missing
    package or an empty extraction falls back to the BeautifulSoup tag strip.
    Either way the trailing reference/citation sections are cut (see above).
    """
    text = ""
    try:
        import trafilatura
        text = trafilatura.extract(raw, include_comments=False, include_tables=True) or ""
    except Exception:  # noqa: BLE001 - not installed / extraction failure -> fallback
        text = ""
    if not text.strip():
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "form"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    return _strip_trailing_sections(text)


def _extract_html(path: Path, encoding: str) -> str:
    raw = _decode(_read_bytes(path), encoding)
    return _clean_html(raw)


# Optional Docling structure-aware ingestion (layout + tables + OCR -> markdown).
# Heavy (torch models) and opt-in; everything fail-soft so a missing package or a
# failed conversion falls back to the lightweight extractors above.
_DOCLING_EXTS = {".pdf", ".docx", ".pptx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}
_DOCLING_CONVERTER = None  # cached; DocumentConverter is expensive to build


def _docling_converter():
    global _DOCLING_CONVERTER
    if _DOCLING_CONVERTER is None:
        # Windows blocks the symlinks HuggingFace uses to populate its model
        # cache unless Developer Mode/admin is on, raising WinError 1314 mid
        # download. Force copies instead. Must be set before huggingface_hub is
        # imported; ingestion runs before the GLiNER2/HF import, so this holds.
        import os
        os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS", "1")
        from docling.document_converter import DocumentConverter
        _DOCLING_CONVERTER = DocumentConverter()
    return _DOCLING_CONVERTER


def _extract_docling(path: Path) -> Optional[str]:
    """Convert a document with Docling and return Markdown, or None on failure."""
    try:
        result = _docling_converter().convert(str(path))
        return result.document.export_to_markdown()
    except Exception as exc:  # noqa: BLE001 - fail soft, caller falls back
        logger.warning("Docling failed on %s (%s); using fallback extractor.", path, exc)
        return None


def normalize_text(text: str) -> str:
    """Normalize whitespace and line endings without destroying structure."""
    # Strip hyphenation soft hyphens and repair umlaut mojibake from misread RTF
    # codepages, so GLiNER/spaCy see clean German words ("Kaiserslautern", "Thürling").
    from postprocess.aggregator import _repair_text
    text = _repair_text(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ").replace("​", "")
    lines = [ln.rstrip() for ln in text.split("\n")]
    # Collapse runs of 3+ blank lines down to 2.
    out: list[str] = []
    blanks = 0
    for ln in lines:
        if ln.strip() == "":
            blanks += 1
            if blanks <= 2:
                out.append("")
        else:
            blanks = 0
            out.append(ln)
    return "\n".join(out).strip()


def extract_text(path: Path, encoding: str = "auto", use_docling: bool = False) -> str:
    """Extract plaintext from a single file based on its extension.

    When ``use_docling`` is set and the type is one Docling handles, try Docling
    first (preserves tables/reading order, OCRs scanned PDFs) and fall back to the
    lightweight extractor if it fails or yields nothing.
    """
    suffix = path.suffix.lower()
    if use_docling and suffix in _DOCLING_EXTS:
        md = _extract_docling(path)
        if md and md.strip():
            return normalize_text(md)
        # else: fall through to the standard extractors below
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix == ".rtf":
        text = _extract_rtf(path, encoding)
    elif suffix in {".html", ".htm"}:
        text = _extract_html(path, encoding)
    else:
        text = _decode(_read_bytes(path), encoding)
    return normalize_text(text)


def iter_input_files(input_path: str | Path, glob: str = "**/*",
                     use_docling: bool = False) -> Iterator[Path]:
    """Yield candidate input files from a file or directory path."""
    p = Path(input_path)
    if p.is_file():
        yield p
        return
    if not p.exists():
        raise FileNotFoundError(f"Input path does not exist: {p}")
    # With Docling on, image/pptx types it can read are no longer "unsupported".
    skip = _SKIP_SUFFIXES - _DOCLING_EXTS if use_docling else _SKIP_SUFFIXES
    for child in sorted(p.glob(glob)):
        if not child.is_file():
            continue
        if child.suffix.lower() in skip:
            logger.debug("Skipping unsupported binary file: %s", child)
            continue
        yield child


def load_documents(input_path: str | Path, glob: str = "**/*",
                   encoding: str = "auto", use_docling: bool = False) -> list[Document]:
    """Ingest all input files into :class:`Document` objects.

    Files that raise during extraction are skipped with a warning so a single
    corrupt file never aborts the run.
    """
    docs: list[Document] = []
    for fp in iter_input_files(input_path, glob, use_docling=use_docling):
        try:
            text = extract_text(fp, encoding, use_docling=use_docling)
        except Exception as exc:  # noqa: BLE001 - fail soft per file
            logger.warning("Failed to extract %s: %s", fp, exc)
            continue
        if not text.strip():
            logger.debug("Empty after extraction, skipping: %s", fp)
            continue
        doc_id = stable_id(str(fp.resolve()), prefix="doc_", length=10)
        docs.append(
            Document(
                doc_id=doc_id,
                source_path=str(fp),
                text=text,
                meta={"filename": fp.name, "suffix": fp.suffix.lower(),
                      "n_chars": len(text)},
            )
        )
    logger.info("Loaded %d documents from %s", len(docs), input_path)
    return docs


# Web / URL ingestion
_USER_AGENT = "SNA-Extraction-Pipeline/1.0 (academic research)"


def is_url(s: str) -> bool:
    """True if ``s`` looks like an http(s) URL."""
    return isinstance(s, str) and s.strip().lower().startswith(("http://", "https://"))


def fetch_url(url: str, timeout: int = 30) -> str:
    """Fetch a single URL and return readable plaintext.

    Handles HTML (boilerplate stripped) and PDF (``application/pdf``) responses;
    other content types are treated as plaintext. Raises on network/HTTP errors
    so the caller can fail soft per URL.
    """
    import requests
    resp = requests.get(url, headers={"User-Agent": _USER_AGENT}, timeout=timeout)
    resp.raise_for_status()
    ctype = (resp.headers.get("Content-Type") or "").lower()

    # requests falls back to ISO-8859-1 when the server sends no charset, which
    # mangles UTF-8 pages. Re-sniff from the body in that case before resp.text.
    if (resp.encoding or "").lower() == "iso-8859-1" and "charset" not in ctype:
        resp.encoding = resp.apparent_encoding or resp.encoding

    if "pdf" in ctype or url.lower().endswith(".pdf"):
        import fitz  # PyMuPDF
        parts: list[str] = []
        with fitz.open(stream=resp.content, filetype="pdf") as doc:
            for page in doc:
                parts.append(page.get_text("text"))
        text = "\n".join(parts)
    elif "html" in ctype or resp.text.lstrip()[:1] == "<":
        text = _clean_html(resp.text)
    else:
        text = resp.text
    return normalize_text(text)


def documents_from_urls(urls: list[str], timeout: int = 30) -> list[Document]:
    """Fetch a list of URLs into :class:`Document` objects (fail-soft per URL)."""
    docs: list[Document] = []
    for url in urls:
        url = url.strip()
        if not url or url.startswith("#"):
            continue
        try:
            text = fetch_url(url, timeout=timeout)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to fetch %s: %s", url, exc)
            continue
        if not text.strip():
            logger.debug("Empty after fetch, skipping: %s", url)
            continue
        docs.append(
            Document(
                doc_id=stable_id(url, prefix="url_", length=10),
                source_path=url,
                text=text,
                meta={"filename": url, "source_type": "url", "n_chars": len(text)},
            )
        )
    logger.info("Fetched %d documents from %d URLs", len(docs), len(urls))
    return docs


def document_from_text(text: str, name: str = "direct_input") -> Document:
    """Wrap a raw text string as a single :class:`Document` (e.g. pasted input)."""
    text = normalize_text(text)
    return Document(
        doc_id=stable_id(name, text[:200], prefix="txt_", length=10),
        source_path=name,
        text=text,
        meta={"filename": name, "source_type": "text", "n_chars": len(text)},
    )


def write_documents_snapshot(documents: list[Document], path: str | Path) -> int:
    """Write the gathered corpus to a portable JSONL snapshot (one Document/line).

    This is the ingestion checkpoint: crawl + fetch + RTF/mojibake repair +
    main-content extraction done once, frozen to a single file. Ship it to another
    machine (the 5090) or re-run extraction in any --mode without re-scraping. The
    doc_ids are preserved, so a snapshot run and a live run produce the same nodes.
    Returns the count written."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        for d in documents:
            fh.write(json.dumps(d.to_dict(), ensure_ascii=False) + "\n")
    return len(documents)


def read_documents_snapshot(path: str | Path) -> list[Document]:
    """Load a corpus JSONL written by write_documents_snapshot. The portable,
    mode-independent input - no crawl, no fetch, no file walk. Bad lines are
    skipped, not raised, so a partially-copied file still loads what it can."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Documents snapshot not found: {path}")
    out: list[Document] = []
    for ln in path.read_text(encoding="utf-8").splitlines():
        ln = ln.strip()
        if not ln:
            continue
        try:
            out.append(Document.from_dict(json.loads(ln)))
        except (ValueError, TypeError) as exc:  # malformed/truncated line
            logger.warning("Skipping bad snapshot line: %s", exc)
    return out


def read_urls_file(path: str | Path) -> list[str]:
    """Read a newline-delimited list of URLs (``#`` comments allowed)."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"URLs file not found: {p}")
    return [ln.strip() for ln in p.read_text(encoding="utf-8").splitlines()
            if ln.strip() and not ln.strip().startswith("#")]


# Unified collector (files + URLs + direct text)
def gather_documents(
    input_path: str | Path | None = None,
    glob: str = "**/*",
    encoding: str = "auto",
    urls: Optional[list[str]] = None,
    urls_file: str | Path | None = None,
    text: Optional[str] = None,
    timeout: int = 30,
    use_docling: bool = False,
) -> list[Document]:
    """Collect documents from every supported source into one list.

    Sources (any combination):
      * ``input_path`` - a file, a directory (globbed), **or an http(s) URL**.
      * ``urls`` / ``urls_file`` - web pages / PDFs fetched over HTTP.
      * ``text`` - a raw string (direct/pasted input).

    Duplicate ``doc_id`` values are removed (first occurrence wins).
    """
    docs: list[Document] = []
    all_urls: list[str] = list(urls or [])

    if input_path:
        if is_url(str(input_path)):
            all_urls.append(str(input_path))
        else:
            docs.extend(load_documents(input_path, glob, encoding, use_docling=use_docling))
    if urls_file:
        all_urls.extend(read_urls_file(urls_file))
    if all_urls:
        docs.extend(documents_from_urls(all_urls, timeout=timeout))
    if text:
        docs.append(document_from_text(text))

    # De-duplicate by doc_id.
    seen: set[str] = set()
    unique: list[Document] = []
    for d in docs:
        if d.doc_id in seen:
            continue
        seen.add(d.doc_id)
        unique.append(d)
    return unique
